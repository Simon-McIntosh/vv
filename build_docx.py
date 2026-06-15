"""
Build the Word version of the ITER VV lateral-displacement report.

This is a structural replica of docs/vv-lateral-displacement-analysis.html. Tables use
a minimal horizontal-rule style (line above and below the header, and a closing rule at
the bottom — no row shading, no vertical lines). Forced-excursion GIFs are placed
side-by-side (no strip images). All section content, callouts and figure captions mirror
the HTML.

    uv run python build_docx.py     ->  docs/vv-lateral-displacement-analysis.docx
"""
from __future__ import annotations
import os
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
DOCS = os.path.join(HERE, "docs")

# palette (matches the HTML CSS)
BLUE      = RGBColor(0x2b, 0x57, 0x97)
DARKBLUE  = RGBColor(0x1a, 0x3a, 0x6e)
GREY      = RGBColor(0x55, 0x55, 0x55)
WHITE     = RGBColor(0xff, 0xff, 0xff)
TABLE_RULE = "888888"        # gray for the horizontal table rules
CALLOUT = {                  # (left-border hex, fill hex) — matches the HTML
    "ok":   ("2a8a2a", "edf8ed"),
    "warn": ("cc8800", "fff8e0"),
    "info": ("2255aa", "e8efff"),
    "key":  ("cc2200", "fff0ee"),
}


# ── low-level OOXML helpers ─────────────────────────────────────────────────
def _shade(cell, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set per-edge borders. Each arg is (sz, color_hex) or None (= nil)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, spec in [("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            sz, color = spec
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _set_borders_legacy(cell, edges, color="cccccc", sz=4):
    """Add borders to specific edges (preserves any existing nil settings on others)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)


# ── content builders ────────────────────────────────────────────────────────
def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE
    # blue bottom border on the heading paragraph
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "2b5797")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def callout(doc, kind, runs, title=None):
    """Single-cell shaded table with a coloured left border — matches HTML callouts."""
    border_hex, fill_hex = CALLOUT[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _shade(cell, fill_hex)
    _set_borders_legacy(cell, ["left"], color=border_hex, sz=24)
    _set_borders_legacy(cell, ["top", "bottom", "right"], color=fill_hex, sz=4)
    p = cell.paragraphs[0]
    _no_space(p)
    if title:
        tr = p.add_run(title + "  ")
        tr.bold = True
        tr.font.color.rgb = RGBColor.from_string(border_hex.upper())
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def body(doc, runs):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10.5)
    return p


def pre(doc, text):
    """Monospace block (matches HTML <pre>)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def table(doc, headers, rows):
    """Minimal-rule table: horizontal lines above/below the header and at the bottom only.
    No row shading, no vertical lines, no inter-row separators."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row — top + bottom horizontal rules, nothing else
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        _set_cell_borders(hdr[j], top=(8, TABLE_RULE), bottom=(6, TABLE_RULE),
                          left=None, right=None)
        p = hdr[j].paragraphs[0]; _no_space(p)
        r = p.add_run(h); r.bold = True; r.font.color.rgb = DARKBLUE; r.font.size = Pt(10)

    n = len(rows)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        is_last = (i == n - 1)
        for j, val in enumerate(row):
            _set_cell_borders(
                cells[j],
                top=None,
                bottom=((8, TABLE_RULE) if is_last else None),
                left=None,
                right=None,
            )
            p = cells[j].paragraphs[0]; _no_space(p)
            r = p.add_run(str(val)); r.font.size = Pt(10)
            if j == 0:
                r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def figure(doc, img, caption, width=6.4):
    if not os.path.exists(img):
        body(doc, [(f"[missing figure: {img}]", False)])
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(caption); cr.italic = True; cr.font.size = Pt(8.5); cr.font.color.rgb = GREY
    c.paragraph_format.space_after = Pt(8)


def side_by_side_figures(doc, items, width_each=3.0):
    """items = [(img_path, caption_runs), ...]; caption_runs = list of (text, bold).
    Renders as a borderless 2-cell table with image + caption stacked in each cell."""
    t = doc.add_table(rows=1, cols=len(items))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (img, cap_runs) in enumerate(items):
        cell = t.cell(0, i)
        _set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(img):
            p.add_run().add_picture(img, width=Inches(width_each))
        c = cell.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for text, bold in cap_runs:
            cr = c.add_run(text)
            cr.italic = True; cr.bold = bold
            cr.font.size = Pt(8.5); cr.font.color.rgb = GREY
        c.paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── assemble the document ───────────────────────────────────────────────────
def main():
    r = np.load(os.path.join(HERE, "data", "rattle_mc_5k.npy"))
    p95, p99, p50 = np.percentile(r, 95), np.percentile(r, 99), np.percentile(r, 50)
    rmax = float(r.max())
    h, e = np.histogram(r, bins=60); mode = 0.5 * (e[h.argmax()] + e[h.argmax() + 1])

    doc = Document()
    doc.styles["Normal"].font.name = "Segoe UI"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ── title ───────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    tr = t.add_run("ITER Vacuum Vessel — Lateral Displacement Analysis")
    tr.bold = True; tr.font.size = Pt(19); tr.font.color.rgb = DARKBLUE
    pPr = t._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom"); b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "18")
    b.set(qn("w:space"), "3"); b.set(qn("w:color"), "2b5797"); pbdr.append(b); pPr.append(pbdr)

    auth = doc.add_paragraph()
    ar = auth.add_run("Simon McIntosh"); ar.bold = True; ar.font.size = Pt(11); ar.font.color.rgb = DARKBLUE
    ar2 = auth.add_run("   ·   2026-05-28"); ar2.font.size = Pt(10); ar2.font.color.rgb = GREY

    sub = doc.add_paragraph()
    sr = sub.add_run("5,000-sample Monte Carlo  ·  scipy HiGHS LP  ·  3-DOF rigid-body model  ·  "
                     "two metrics: departure-from-centre (frictionless) and polytope width "
                     "(stiction walk)  ·  R_s ≈ 10 m, M ≈ 9000 t (ITER_D_6TLUDY)")
    sr.font.size = Pt(8.5); sr.font.color.rgb = GREY

    # ── Summary of findings ─────────────────────────────────────────────────
    callout(doc, "key", [
        ("The vessel rests on nine inward-inclined (15°) gravity supports that form a soft "
         "lateral gravitational pendulum (K ≈ 2.7 kN/mm, T_n ≈ 11 s). Whether the vessel "
         "actually self-centres is decided by friction, and that selects between two bounds:\n\n",
         False),
        ("• Frictionless lower bound — self-centring (§3–§6). ", True),
        ("If the supports slid freely the vessel would return to the gravitational centre at "
         "rest (≈ 0 contribution to the n = 1 shift); only forced excursions would reach the "
         f"departure envelope — nominal 1.55 mm, distribution-dependent tail to ~{rmax:.1f} mm.\n",
         False),
        ("• Realistic bound — stiction walk (§10). ", True),
        ("Dead-weight axial load per VVGS ≈ 10 MN; even with low-friction hinges the toroidal "
         "breakaway is ≈ 1 MN per support (≈ 5.7 MN for the whole vessel) — roughly 1000× the "
         "≈ 4 kN gravitational restoring. Gravity cannot recentre the vessel: once a disruption/"
         "VDE load slides it, it stays, and over many events it ratchets within the polytope. "
         "The relevant n = 1 envelope is the polytope width ≈ 3.1 mm (§9) — about half the 6 mm "
         "budget.\n", False),
        ("• Thermal cycling does not recentre it (§11). ", True),
        ("Warming to 100 °C grows the vessel radially ~12 mm (freely accommodated), but the "
         "off-centre toroidal demand is only ~µm — far too small to break the MN-scale stiction.\n",
         False),
        ("• The statistics are only as good as the gap distribution (§12). ", True),
        ("The Uniform(±1.5 mm) prior is not validated, so the MC percentiles are illustrative. "
         "The robust results are the distribution-free bounds and the force comparison. "
         "Measuring the nine gaps with feeler gauges collapses the problem to a deterministic "
         "calculation and is the recommended path.", False),
    ], title="Summary of findings:")

    # ── §1 — Problem & Assumptions ──────────────────────────────────────────
    heading2(doc, "1 — Problem & Assumptions")
    body(doc, [
        ("The vacuum vessel rests on ", False),
        ("nine gravity supports (VVGS)", True),
        (" equally spaced toroidally beneath the lower ports, on a ring of radius R_s ≈ 10 m. "
         "Each VVGS is a ", False),
        ("dual-hinge mechanism inclined at 15° from vertical", True),
        (", leaning inward toward the machine axis. Radial motion of the vessel is permitted "
         "by rotation of the dowels (this accommodates thermal expansion between assembly "
         "and operation/baking); toroidal and vertical motion are restrained by the dowels' "
         "rigidity. The toroidal restraint is a ", False),
        ("±1.5 mm (3 mm total) dowel slot", True),
        (".", False),
    ])
    table(doc, ["Parameter", "Value", "Note"], [
        ["Number of supports", "9", "Equally spaced toroidally"],
        ["Support ring radius R_s", "≈ 10 m", "ITER_D_6TLUDY (VV Load Spec)"],
        ["Toroidal slot per support", "±1.5 mm (3 mm total)", "Assembly tolerance (see below)"],
        ["Hinge inclination from vertical", "15°", "Inward (centring)"],
        ["Supported mass M", "≈ 9000 t", "VV + in-vessel (ITER_D_6TLUDY)"],
    ])

    heading3(doc, "Assembly assumption — the source of mobility")
    body(doc, [
        ("At each VVGS the connecting pin is oriented in the toroidal direction and runs "
         "through interleaved vessel-side and ground-side tabs. ", False),
        ("The holes are tight on the pin (no radial play); the toroidal clearance is the "
         "spacing between adjacent tabs along the pin.", True),
        (" That inter-tab spacing is required to bring the parts together so the tabs can "
         "be aligned and the pin can be inserted — without it the support cannot be "
         "assembled. The same spacing then permits the vessel-side tabs to slide along the "
         "pin (toroidally) once assembled, giving ", False),
        ("±1.5 mm (3 mm total)", True),
        (" of lateral mobility per support. If the joints are ", False),
        ("left unshimmed", True),
        (" — the case treated throughout this report — this assembly clearance is what "
         "drives the analysis.", False),
    ])
    body(doc, [
        ("For the Monte-Carlo sections (§5–§6) we adopt the prior that each support's pin "
         "offset u_i is ", False),
        ("independent across the 9 supports", True),
        (" and ", False),
        ("uniformly distributed on [−1.5 mm, +1.5 mm]", True),
        (". This prior is not validated", True),
        (" — neither the uniform shape nor the ±1.5 mm range is established (§12) — so the "
         "resulting percentiles are illustrative, and the load-bearing conclusions of this "
         "report are deliberately the distribution-free ones (the kinematic bounds and the "
         "force comparison). Shimming a support fixes its u_i to a measured value and removes "
         "that support's contribution to the mobility (§6).", False),
    ])

    heading3(doc, "Gravitational centring")
    body(doc, [
        ("The inward 15° inclination is deliberate. The nine support axes converge above the "
         "vessel centre of gravity; for the lateral (n = 1) mode the vessel therefore "
         "behaves as if suspended from that convergence point — an ", False),
        ("inclined-hinge gravitational pendulum", True),
        (" — and the lateral plane has a parabolic potential well centred on the nominal "
         "position. The effective pendulum length is L_eff ≈ R_s / tan(15°) − z_CoG ≈ 32 m "
         "(≈ 26–32 m over the estimated z_CoG = 5–10 m range).", False),
    ])
    table(doc, ["Quantity", "Value"], [
        ["Lateral centring stiffness", "K ≈ 2.7 kN / mm"],
        ["Natural period", "T_n ≈ 11 s   (f ≈ 0.09 Hz)"],
        ["Gravitational force to push the vessel to a ±1.5 mm stop", "F_stop ≈ 4.1 kN"],
        ["Rest position with no lateral force (frictionless)", "q* = 0 (centred)"],
    ])
    callout(doc, "info", [
        ("With no friction, and no applied lateral force, the vessel would sit at the "
         "gravitational centre regardless of the assembly offsets — this is the lower-bound "
         "picture of §3–§6. §10 shows that dead-weight friction at the supports is ~1000× the "
         "centring force, so in reality the vessel does not return to centre once displaced; it "
         "holds its last position. Both bounds are carried through this report.", False),
    ], title="Consequence — in the frictionless idealisation.")
    figure(doc, os.path.join(DOCS, "diagrams", "D_rocking_pendulum.gif"),
           "The mechanism, side-on. Each VVGS is its actual 15° inclined dual-hinge "
           "(parallelogram) 4-bar; extended, the strut axes converge at the virtual pivot P, "
           "and the rigid vessel rocks about P as a soft gravitational pendulum (lateral sway "
           "exaggerated ~×400). This restoring mechanism is ~1000× too weak to overcome the "
           "supports' dead-weight stiction (§10).", width=3.6)

    # ── §2 — Constraint Model, Polytope & Linear Programming ────────────────
    heading2(doc, "2 — Constraint Model, Polytope & Linear Programming")
    body(doc, [
        ("We model the vacuum vessel as a rigid body with three in-plane degrees of "
         "freedom — two horizontal translations Δx, Δy and a rotation Δθ about the vertical "
         "axis (collected as ", False),
        ("q = [Δx, Δy, Δθ]", True),
        ("). Each of the nine supports imposes a single linear inequality constraint: the "
         "toroidal slide of the vessel relative to its dowel must stay within the ±1.5 mm "
         "slot. For support i at toroidal angle φ_i = π/2 − 2πi/9 on the support ring at "
         "radius R = 10 m, the slide is", False),
    ])
    pre(doc,
        "  δᵢ = −sin(φᵢ)·Δx + cos(φᵢ)·Δy + R·Δθ   ≡   A[i, :] · q ,\n"
        "      |uᵢ + δᵢ| ≤ 1.5 mm     for all i = 1 … 9,")
    body(doc, [
        ("so the displacement is bounded by ", False),
        ("18 linear inequalities", True),
        (" (one upper and one lower bound per support). The 9×3 constraint matrix has ", False),
        ("AᵀA = diag(4.5, 4.5, 900)", True),
        (" — exactly diagonal, so the three rigid-body DOFs are mutually uncorrelated.", False),
    ])
    callout(doc, "info", [
        ("The set of vessel positions that simultaneously satisfy all 18 inequalities is a "
         "convex region of (Δx, Δy, Δθ) space bounded by flat faces — geometrically a ", False),
        ("polytope", True),
        (" (the three-dimensional generalisation of a polygon). Projected onto the horizontal "
         "plane (Δx, Δy), with Δθ optimised at each direction, it is a closed convex polygon "
         "of feasible lateral VV-centre positions — the ", False),
        ("displacement polytope", True),
        (". Its shape depends on the assembly offsets and its diameter is set by the "
         "±1.5 mm slot.\n\n", False),
        ("A ", False),
        ("linear program (LP)", True),
        (" optimises a linear objective subject to linear inequality constraints — the "
         "canonical method for problems of exactly this kind. We solve, in each direction θ, "
         "the LP that maximises the projection of q on the unit vector (cos θ, sin θ, 0) "
         "subject to the 18 inequalities above; the optimum lies at the polytope vertex "
         "furthest along that direction. Sweeping θ traces the polytope boundary, and the "
         "maximum of ||q[:2]|| over the sweep is our metric: the ", False),
        ("maximum lateral displacement of the VV centre from the gravitational centre", True),
        (" that any applied force can drive the vessel to. We use the HiGHS solver "
         "(scipy.optimize.linprog); each LP solves in milliseconds.", False),
    ], title="Polytope & LP.")

    # ── §3 — State Diagrams ─────────────────────────────────────────────────
    heading2(doc, "3 — State Diagrams")
    body(doc, [
        ("Each panel shows the VV ring at its maximum-displacement position; the red arrow "
         "from the machine axis (black cross) to the displaced VV centre is the "
         "maximum-displacement vector — its direction and length show where and how far the "
         "vessel can be driven by an applied lateral force. The blue shaded polygon around "
         "the machine axis is the displacement polytope (the kinematic envelope of all "
         "reachable VV-centre positions). Displacements are magnified ×500.", False),
    ])
    figure(doc, os.path.join(DOCS, "plots", "vv_states.png"),
           "Left — nominal assembly (u = 0): the polytope is symmetric about the machine "
           "axis and the maximum displacement is 1.55 mm in any direction. Right — worst "
           "observed MC sample: the assembly offsets shift the polytope off-axis, so the "
           "maximum-displacement vector reaches further from the machine axis (2.98 mm). At "
           "each support the toroidal slot (grey, with red ±1.5 mm end-stops) is held at a "
           "constant radial distance from the vessel wall by the inclined linkage — radial "
           "motion is free; the coloured pin in the slot indicates how much of the ±1.5 mm "
           "toroidal travel is taken up (blue = slack, orange = near the limit, red = at the "
           "stop).")

    # ── §4 — Forced-Excursion Envelope ──────────────────────────────────────
    heading2(doc, "4 — Forced-Excursion Envelope")
    body(doc, [
        ("The animations sweep the VV centre through the polytope along its principal "
         "axis — the kinematic envelope of possible lateral motion about the centred rest "
         "position. With no applied force the vessel sits at the machine axis; sufficient "
         "applied force drives it through the envelope. ×500 magnification; the grey "
         "toroidal slot at each support tracks the vessel wall (radial motion is free), and "
         "the coloured pin shows the toroidal-constraint usage (blue = slack, orange = near "
         "limit, red = at stop).", False),
    ])
    side_by_side_figures(doc, [
        (os.path.join(DOCS, "animations", "rattle_worst_case.gif"),
         [("Worst observed MC sample. ", True),
          (f"Offset assembly; maximum forced departure from the machine axis ≈ "
           f"{rmax:.2f} mm — the kinematic ceiling for the worst random assembly drawn from "
           "the §1 uniform prior.", False)]),
        (os.path.join(DOCS, "animations", "rattle_mode.gif"),
         [("Typical as-built assembly ", True),
          (f"(drawn near the distribution mode). Maximum forced departure ≈ {mode:.2f} mm — "
           "the most likely value an as-built machine will actually exhibit.", False)]),
    ], width_each=3.0)

    # ── §5 — Monte Carlo Distribution ───────────────────────────────────────
    heading2(doc, "5 — Monte Carlo Distribution")
    body(doc, [
        ("Five thousand assemblies, with each support's offset drawn independently from "
         "Uniform(−1.5 mm, +1.5 mm). For each assembly the maximum lateral displacement of "
         "the VV centre from the gravitational centre is computed (72-direction LP). "
         "Reproducibly generated by vv_mc_generator.py with fixed seed 20260527.", False),
    ])
    figure(doc, os.path.join(DOCS, "plots", "mc_dashboard.png"),
           "Distribution (left) and CDF (right) of the maximum lateral displacement of the "
           "VV centre from the gravitational centre, over 5000 randomly-drawn assemblies. "
           "The nominal (centred-assembly) value is 1.55 mm; offset assemblies can reach "
           "further from the centre because the polytope is shifted off-axis.")
    table(doc, ["Statistic", "Max lateral displacement (mm)"], [
        ["Nominal (u = 0) — symmetric envelope half-width", "1.55"],
        ["Mode", f"{mode:.2f}"],
        ["Median (P50)", f"{p50:.2f}"],
        ["P90", "2.12"],
        ["P95", f"{p95:.2f}"],
        ["P99", f"{p99:.2f}"],
        ["Max observed (n = 5000)", f"{rmax:.2f}"],
    ])
    callout(doc, "warn", [
        ("The Monte-Carlo percentiles above are only as trustworthy as the assumed input "
         "distribution. We drew each support offset from Uniform(−1.5 mm, +1.5 mm) "
         "independently, but neither the uniform shape nor the ±1.5 mm range is validated — "
         "the as-built gaps are set by assembly history, not by a known random process. The "
         "P95/P99 figures therefore carry low confidence and are illustrative. What does not "
         "depend on the distribution are (i) the nominal centred envelope (1.55 mm), (ii) the "
         "polytope-width ceiling (§9), and (iii) the stiction-vs-centring force comparison "
         "(§10). The recommendation (§12) is to measure the nine gaps with feeler gauges "
         "rather than propagate an unvalidated prior.", False),
    ], title="Reliability caveat — read the percentiles with caution.")

    # ── §6 — Effect of Gap Measurements ─────────────────────────────────────
    heading2(doc, "6 — Effect of Gap Measurements")
    body(doc, [
        ("Each measured gap fixes one u_i to its true value. Because the polytope's offset "
         "from the gravitational centre is governed by the assembly offsets, ", False),
        ("every measured gap directly tightens the conditional displacement bound", True),
        (": each measurement pulls the polytope toward the centre. With all 9 measured and "
         "found near-centred, the conditional maximum displacement collapses to the 1.55 mm "
         "nominal envelope.", False),
    ])
    figure(doc, os.path.join(DOCS, "plots", "partial_measurement.png"),
           "Left — distribution of the maximum lateral displacement; measuring all 9 gaps "
           "reveals one value from this distribution. Right — conditional P95 vs number of "
           "measured sectors, for two scenarios of the measured values: centred (lowest "
           "residual bound) and typical random (mid-range). Because AᵀA is exactly diagonal "
           "the three rigid-body DOFs are uncorrelated, so only the number of measured "
           "sectors matters; adjacent vs spread-out is statistically identical.")
    table(doc, ["Measured sectors (k of 9)", "P95 — centred (mm)", "P95 — typical (mm)"], [
        ["0 (baseline)", f"{p95:.2f}", f"{p95:.2f}"],
        ["1", "≈ 2.26", "≈ 2.26"],
        ["3", "≈ 2.19", "≈ 2.22"],
        ["5", "≈ 2.05", "≈ 2.00"],
        ["9 (all)", "→ 1.55 (nominal env.)", "deterministic"],
    ])

    # ── §7 — Start-Up Lateral Loads: Passive Self-Alignment ─────────────────
    heading2(doc, "7 — Start-Up Lateral Loads: Passive Self-Alignment")
    body(doc, [
        ("During plasma current ramp-up (0 → ~3 MA over ~10 s), the time-changing plasma "
         "current induces image currents in the vessel. These currents interact with the "
         "toroidal-field asymmetry produced by TF vault closure tolerances (a "
         "long-wavelength n ≤ 4 perturbation) to produce a lateral n = 1 force on the "
         "vessel.", False),
    ])
    callout(doc, "key", [
        ("The lateral force on the induced vessel current is such that the vessel is "
         "pulled ", False),
        ("toward", True),
        (" the n = 1 magnetic axis of the (offset) toroidal field — passive vessel currents "
         "act to ", False),
        ("self-align", True),
        (" the wall with the field profile. ", False),
        ("Any vessel motion during the ramp therefore reduces the peak-to-peak misalignment "
         "between the n ≤ 4 first-wall and toroidal-field profiles", True),
        ("; it consumes none of the 6 mm budget, and in fact frees some of it back. The "
         "start-up ramp is a budget-relieving event for the n = 1 first-wall criterion, not "
         "a budget-eroding one.", False),
    ], title="Important: the start-up displacement direction is favourable.")

    heading3(doc, "Force scale & timescale")
    body(doc, [
        ("Order-of-magnitude: with vessel image current of O(10⁵ A) during the ramp, "
         "B_T ≈ 5 T, fractional n = 1 field error ε ~ 10⁻⁴–10⁻³, and a toroidal length scale "
         "~ 2πR, the lateral force F_ramp ~ ε · I_vv · B_T · L is of order ", False),
        ("1–10 kN", True),
        (". This is comparable to the gravitational centring threshold F_stop ≈ 4.1 kN but is "
         "far below the ≈ 1 MN per-support toroidal breakaway force (§10): a kN-scale start-up "
         "force does not slide the supports at all. The ramp duration (10 s) is comparable to "
         "the pendulum natural period (T_n ≈ 11 s), so in the frictionless idealisation the "
         "quasi-static deflection F/K = 0.4–4 mm would appear as a peak transient excursion of "
         "order ", False),
        ("1–4 mm", True),
        (" in the favourable (self-aligning) direction. In the realistic stiction-dominated "
         "picture the start-up force is orders of magnitude below breakaway, so it produces no "
         "net lateral motion — it neither erodes nor relieves the budget kinematically. Either "
         "way this source is not budget-limiting; a detailed EM transient is required to pin "
         "the precise value.", False),
    ])

    # ── §8 — Disruption Loads & VVGS Impact Case ───────────────────────────
    heading2(doc, "8 — Disruption Loads & VVGS Impact Case")
    body(doc, [
        ("Disruptions and asymmetric vertical displacement events (VDEs) deliver the large "
         "lateral loads. ", False),
        ("The bounding magnitudes and durations are defined in the ITER VV Load Specification "
         "(VVLS)", True),
        (" and are substantially larger than the illustrative impulse used in an earlier draft "
         "of this note; the values below are order-of-magnitude only and should be taken from "
         "the VVLS for design.", False),
    ])
    callout(doc, "warn", [
        ("An earlier version quoted a lateral impulse J ~ 10⁵ N·s over ~10 ms (free-mass "
         "velocity ~12 mm/s). Both the magnitude and the duration are lower than the VVLS "
         "bounding values, so the resulting toroidal VVGS force was correspondingly "
         "under-stated. The VVLS net horizontal VV loads are of order MN to tens of MN.", False),
    ], title="Correction.")
    body(doc, [
        ("Two regimes follow from the friction analysis (§10), with breakaway ≈ 5.7 MN "
         "aggregate (≈ 1 MN per support):", False),
    ])
    body(doc, [
        ("• Small / illustrative impulse (below breakaway). ", True),
        ("The earlier J ~ 10⁵ N·s carries kinetic energy ½Mv² ≈ 0.56 kJ (M ≈ 9000 t). The "
         "friction work over a full 3 mm slide is ≈ 5.7 MN × 3 mm ≈ 17 kJ ≫ 0.56 kJ, so such "
         "an impulse would move the vessel only ~0.1 mm before stiction arrests it — there is "
         "no free-flight impact on the stops, contrary to the earlier free-mass picture.", False),
    ])
    body(doc, [
        ("• VVLS bounding loads (above breakaway). ", True),
        ("Net horizontal loads of MN-to-tens-of-MN exceed the ≈ 5.7 MN breakaway, so the "
         "vessel does slide during major events. Each event displaces it by an amount set by "
         "the load, the impulse and the remaining slot; the displacement is retained (gravity "
         "cannot pull it back, §10). Repeated events make the vessel ratchet within the "
         "polytope (§9–§10).", False),
    ])
    callout(doc, "key", [
        ("— and it is the mechanism that drives the vessel away from the gravitational centre "
         "toward the polytope-width bound. Its magnitude must be read from the VVLS; the "
         "analysis here establishes the consequence (irreversible walk), not the load value.",
         False),
    ], title="The disruption/VDE toroidal force is the bounding VVGS lateral load case")

    # ── §9 — Maximum Rattle: the Polytope-Width Bound ──────────────────────
    heading2(doc, "9 — Maximum Rattle: the Polytope-Width Bound")
    body(doc, [
        ("§3–§6 measured the departure from the gravitational centre — the right metric if "
         "the vessel self-centres. This section measures the complementary quantity: the ",
         False),
        ("full peak-to-peak width of the displacement polytope", True),
        (" (its diameter along the worst direction). The width is the total lateral travel "
         "available to the vessel — the kinematic ceiling on how far it can walk if friction "
         "(not gravity) sets its rest position (§10).", False),
    ])
    callout(doc, "info", [
        ("Departure-from-centre is minimised at the centred assembly (u = 0 → 1.55 mm) and "
         "grows for offset assemblies. The polytope width is maximised at the centred assembly "
         "and shrinks for offset assemblies, so the worst-case width is the nominal one. "
         "Nominal width = 3.09 mm (distribution-free ceiling); over the assembly population: "
         "median 2.24 mm, P5 1.19 mm, minimum 0.21 mm — a typical as-built assembly has a "
         "narrower reachable envelope than nominal.", False),
    ], title="Width and departure behave oppositely.")
    figure(doc, os.path.join(DOCS, "plots", "rattle_dashboard.png"),
           "Distribution (left) and CDF (right) of the polytope width (peak-to-peak rattle) "
           "over 5000 assemblies. The nominal (u = 0) width 3.09 mm is the ceiling; offset "
           "assemblies are more constrained and sit below it. The shape inherits the "
           "unvalidated gap prior (§12), but the ceiling is distribution-free.")
    figure(doc, os.path.join(DOCS, "animations", "rattle_width_nominal.gif"),
           "The nominal assembly swept across the full polytope width (×500). The VV centre "
           "travels the complete ±1.55 mm diameter; the coloured pins show the toroidal-slot "
           "usage (blue = slack, orange = near limit, red = at the ±1.5 mm stop).", width=3.6)

    # ── §10 — Which Bound? ──────────────────────────────────────────────────
    heading2(doc, "10 — Which Bound? Frictionless Self-Centring vs Stiction Walk")
    body(doc, [
        ("Which regime ITER occupies is decided by one comparison: the gravitational "
         "restoring force that would recentre the vessel, versus the friction (stiction) "
         "force that resists toroidal sliding at the supports.", False),
    ])
    table(doc, ["Quantity", "Value", "Basis"], [
        ["Dead-weight axial force per VVGS (15° strut)", "≈ 10.2 MN", "(Mg/9)/cos 15°, M ≈ 9000 t"],
        ["Toroidal breakaway (stiction) per support", "≈ 1.0 MN", "μ ≈ 0.1 × axial"],
        ["Aggregate lateral breakaway (all 9, worst dir.)", "≈ 5.7 MN", "Σ μN |sin(θ−φ_i)|"],
        ["Gravitational restoring at the ±1.5 mm stop", "≈ 4.1 kN", "K × 1.5 mm"],
        ["Breakaway ÷ restoring", "≈ 1400×", "2–3 orders of magnitude"],
    ])
    callout(doc, "key", [
        ("The gravitational restoring force (kN) is negligible against the toroidal breakaway "
         "force (MN). The vessel does not self-centre. Once any load large enough to break "
         "stiction (a disruption/VDE, §8) slides the supports, the vessel moves and stays "
         "there — gravity cannot pull it back. Over many events the vessel ratchets / walks "
         "through the polytope. The frictionless self-centring picture of §3–§6 is a lower "
         "bound; the realistic operational envelope is the polytope width (§9), ≈ 3.1 mm at "
         "worst.", False),
    ], title="Stiction wins by ~1000×.")
    figure(doc, os.path.join(DOCS, "plots", "two_bounds.png"),
           "Left — the two metrics as distributions: departure-from-centre (blue, frictionless "
           "lower bound) and polytope width (red, stiction-walk envelope). Right — the VV n = 1 "
           "first-wall contribution under each regime against the 6 mm budget: frictionless "
           "self-centred ≈ 0, but the realistic stiction-walk ceiling ≈ 3.1 mm consumes about "
           "half the budget.")

    # ── §11 — Thermal Cycling & the Overdetermined Structure ───────────────
    heading2(doc, "11 — Thermal Cycling & the Statically-Overdetermined Structure")
    body(doc, [
        ("Does warming the vessel relax the nine stictions in a coordinated way that lets it "
         "settle back to the gravitational centre? And what loads does an off-centre vessel "
         "generate when heated, given that nine supports over-constrain a body with only "
         "three in-plane freedoms?", False),
    ])
    heading3(doc, "Free radial breathing — large, but it does not recentre")
    body(doc, [
        ("For ΔT ≈ 78 K (22 → 100 °C) and 316L(N) α ≈ 16×10⁻⁶ K⁻¹, the support radius grows "
         "by R_s·α·ΔT ≈ 12.5 mm (≈ 32 mm at the full 200 °C bake). This is the motion the "
         "dual-hinge supports are designed to take: the dowels rotate, the radial growth is "
         "free, and no toroidal slot demand or lateral load is generated — provided the "
         "expansion is about the support-ring centre.", False),
    ])
    callout(doc, "key", [
        ("If the vessel is off-centre by d, isotropic expansion about its centroid produces a "
         "toroidal slot demand of only α·ΔT·|d| ≈ 3.7 µm (for d = 3 mm). The elastic force to "
         "absorb such a micron-scale mismatch at a locked slot is a few kN — ~1000× below the "
         "≈ 1 MN stiction. Heating neither unloads the supports (weight, hence stiction, is "
         "temperature-independent) nor generates enough toroidal demand to slide them. An "
         "off-centre vessel stays off-centre through thermal cycles; there is no coordinated "
         "relaxation back to the gravitational centre.", False),
    ], title="Thermal cycling does not break the stiction.")
    heading3(doc, "Locked-in loads: uniform expansion is benign; asymmetry is the risk")
    body(doc, [
        ("Uniform isotropic expansion of an off-centre vessel maps exactly onto a "
         "representable rigid-body shift, because δ_i^th = α·ΔT·(d·ê_t,i) = A[i,:]·(α·ΔT·d_x, "
         "α·ΔT·d_y, 0). A demand lying in the column space of A is a rigid-body mode and "
         "induces no self-stress (only the µm-scale take-up above if a slot is locked). So "
         "uniform heating of a nominally symmetric vessel is benign.", False),
    ])
    body(doc, [
        ("The genuine thermal risk is support-to-support asymmetry: non-uniform temperatures "
         "(sector gradients, asymmetric baking, local heating) produce a toroidal demand "
         "pattern that does not lie in the column space of A. That component cannot be relieved "
         "by any rigid-body motion, so in a statically over-determined ring it is carried as a "
         "self-equilibrating internal load set across the redundant supports — locked-in "
         "toroidal dowel forces scaling with the differential expansion (α·ΔT·ΔR_support-to-"
         "support), not with the mean temperature.", False),
    ])
    callout(doc, "warn", [
        ("The dominant thermal lateral concern is not mean-temperature recentring (it does not "
         "happen) but the locked-in self-stress from asymmetric thermal fields acting on a "
         "stiction-locked, over-determined 9-support ring. A thermo-mechanical FE model with "
         "realistic sector temperature spreads and the measured/assumed slot states is required "
         "to size these internal loads and confirm they stay within the VVGS toroidal capacity.",
         False),
    ], title="Conjecture (needs FE confirmation).")

    # ── §12 — Reliability of the Gap Distribution & Measurement ────────────
    heading2(doc, "12 — Reliability of the Gap Distribution & the Case for Measurement")
    body(doc, [
        ("The statistics in §5–§6 rest on one weak assumption: that each support's offset is "
         "an independent draw from Uniform(−1.5 mm, +1.5 mm). Neither part is justified — the "
         "real gaps are the product of an assembly sequence, not a random process, and there "
         "is no evidence the range is ±1.5 mm or the shape uniform. A different but equally "
         "defensible prior moves the percentiles materially, so P95/P99 carry little weight.",
         False),
    ])
    table(doc, ["Result", "Depends on the gap distribution?", "Confidence"], [
        ["Nominal centred envelope (1.55 mm)", "No — geometry + slot", "HIGH"],
        ["Polytope-width ceiling (3.09 mm)", "No — geometry + slot", "HIGH"],
        ["Stiction ≫ centring (~1000×) & walk verdict", "No — force comparison", "HIGH"],
        ["Thermal no-recentring verdict", "No — force comparison", "HIGH"],
        ["MC percentiles (P95 = 2.27 mm, etc.)", "Yes — strongly", "LOW"],
    ])
    callout(doc, "ok", [
        ("Because the conclusions that matter are distribution-free, and the only "
         "distribution-dependent outputs are the unreliable percentiles, the highest-value "
         "action is to measure the nine VVGS gaps directly with feeler gauges. Nine "
         "measurements turn the problem into a deterministic kinematic calculation — the exact "
         "polytope, its width, and the as-built offset of the vessel centre — with no prior to "
         "defend. As the reviewer notes, this is also less effort than building and justifying "
         "an input-uncertainty model. Each measured gap also tightens the conditional bound "
         "(§6); nine collapse it entirely.", False),
    ], title="Recommendation: measure, don't propagate.")

    # ── §13 — The 6 mm n ≤ 4 First-Wall Budget ──────────────────────────────
    heading2(doc, "13 — The 6 mm n ≤ 4 First-Wall Budget")
    body(doc, [
        ("The start-up heat-load criterion limits the peak-to-peak difference between the "
         "n ≤ 4 filtered first-wall profile and the n ≤ 4 filtered toroidal-field profile "
         "to ", False),
        ("≤ 6 mm", True),
        (". The VV n = 1 lateral displacement is one contributor — and the two bounds (§10) "
         "give two very different answers:", False),
    ])
    table(doc, ["Picture", "VV n = 1 contribution to the 6 mm budget"], [
        ["Frictionless self-centring — at rest, between shots", "≈ 0 mm"],
        ["Start-up ramp (passive currents → wall follows TF)", "net reduces mismatch (kN, no slide)"],
        ["From-centre kinematic worst (forced, frictionless)", f"≈ {rmax:.2f} mm"],
        ["Stiction walk — realistic worst case (polytope width)", "≈ 3.1 mm"],
    ])
    callout(doc, "key", [
        ("Because stiction prevents self-centring (§10), the vessel sits wherever it was last "
         "driven, anywhere within the ±-width polytope, until the gaps are measured or the "
         "supports are shimmed. The optimistic ≈ 0 mm (self-centred) value should not be "
         "assumed for the budget. The remaining ~3 mm is what is then available for the "
         "toroidal-field side (mainly TF vault closure n ≤ 4) and other first-wall "
         "contributors — a materially tighter allocation than the self-centred picture implies.",
         False),
    ], title="The realistic VV n = 1 contribution is up to ≈ 3 mm — about half the 6 mm budget — not ≈ 0.")
    callout(doc, "ok", [
        ("Measuring the nine gaps (§12) removes the distribution uncertainty and quantifies "
         "the as-built offset; shimming a support fixes its u_i and shrinks the reachable "
         "polytope (§6), pulling the realistic contribution back toward the nominal envelope. "
         "These are deterministic, low-effort levers — far more reliable than the unvalidated "
         "statistical bound.", False),
    ], title="How to recover budget.")

    # ── footer ──────────────────────────────────────────────────────────────
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Author: Simon McIntosh. Monte-Carlo data generated reproducibly by "
        "vv_mc_generator.py (scipy HiGHS LP, fixed seed 20260527 → committed data/). "
        "Departure-from-centre and polytope-width figures/animations by vv_viz.py and "
        "vv_rattle_figures.py; the stiction/thermal force ledger by vv_mechanics.py; "
        "displacements magnified ×500. Geometry & centring (ITER VV Load Spec + mass table, "
        "ITER_D_6TLUDY): R_s ≈ 10 m, M ≈ 9000 t, 15° inclined dual-hinge → L_eff ≈ 26–32 m → "
        "K ≈ 2.7 kN/mm, T_n ≈ 11 s, F_stop ≈ 4.1 kN; per-VVGS dead-weight axial ≈ 10 MN, "
        "toroidal breakaway ≈ 1 MN. Every number regenerates from the repository. "
        "Source: Simon-McIntosh/vv."
    )
    fr.font.size = Pt(8); fr.font.color.rgb = GREY

    out = os.path.join(DOCS, "vv-lateral-displacement-analysis.docx")
    doc.save(out)
    print(f"-> wrote {out}")


if __name__ == "__main__":
    main()
